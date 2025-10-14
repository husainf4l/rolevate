"""Export services for generating PDF and DOCX files."""
import uuid
from pathlib import Path
from typing import Literal
from loguru import logger

from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

from app.models.schemas import CVData
from app.services.template_filler import TemplateFiller
from app.config import settings


class CVExporter:
    """Export CV data to various formats."""
    
    def __init__(self):
        """Initialize exporter."""
        self.template_filler = TemplateFiller()
        self.output_dir = Path(settings.output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    async def export_to_pdf(
        self,
        cv_data: CVData,
        template_name: str = "modern_cv",
        output_filename: str = None
    ) -> Path:
        """
        Export CV data to PDF.
        
        Args:
            cv_data: Structured CV data
            template_name: Template to use
            output_filename: Custom output filename (optional)
            
        Returns:
            Path: Path to generated PDF file
        """
        logger.info(f"Exporting CV to PDF with template: {template_name}")
        
        try:
            # Render HTML
            html_content = self.template_filler.render_html(cv_data, template_name)
            
            # Get CSS (if available)
            css_content = self.template_filler.get_template_css(template_name)
            
            # Generate filename
            if output_filename is None:
                safe_name = cv_data.full_name.replace(' ', '_').replace('/', '_')
                output_filename = f"{safe_name}_CV_{uuid.uuid4().hex[:8]}.pdf"
            
            if not output_filename.endswith('.pdf'):
                output_filename += '.pdf'
            
            output_path = self.output_dir / output_filename
            
            # Configure fonts
            font_config = FontConfiguration()
            
            # Create PDF
            html_obj = HTML(string=html_content)
            css_obj = CSS(string=css_content, font_config=font_config) if css_content else None
            
            if css_obj:
                html_obj.write_pdf(
                    str(output_path),
                    stylesheets=[css_obj],
                    font_config=font_config
                )
            else:
                html_obj.write_pdf(str(output_path), font_config=font_config)
            
            logger.success(f"PDF exported successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise Exception(f"Failed to export PDF: {e}")
    
    async def export_to_docx(
        self,
        cv_data: CVData,
        template_name: str = "modern_cv",
        output_filename: str = None
    ) -> Path:
        """
        Export CV data to DOCX.
        
        Args:
            cv_data: Structured CV data
            template_name: Template to use
            output_filename: Custom output filename (optional)
            
        Returns:
            Path: Path to generated DOCX file
        """
        logger.info(f"Exporting CV to DOCX")
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
        except ImportError:
            raise ImportError("python-docx is required for DOCX export")
        
        try:
            # Create new document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add header with name and title
            header = doc.add_heading(cv_data.full_name, 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if cv_data.job_title:
                job_title_para = doc.add_paragraph(cv_data.job_title)
                job_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                job_title_para.runs[0].font.size = Pt(14)
                job_title_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            
            # Contact information
            contact_parts = []
            if cv_data.contact.email:
                contact_parts.append(f"✉ {cv_data.contact.email}")
            if cv_data.contact.phone:
                contact_parts.append(f"☎ {cv_data.contact.phone}")
            if cv_data.contact.location:
                contact_parts.append(f"📍 {cv_data.contact.location}")
            
            if contact_parts:
                contact_para = doc.add_paragraph(" | ".join(contact_parts))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_para.runs[0].font.size = Pt(10)
            
            # Professional Summary
            if cv_data.summary:
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(cv_data.summary)
            
            # Experience
            if cv_data.experience:
                doc.add_heading('Professional Experience', level=1)
                for exp in cv_data.experience:
                    # Job title and company
                    job_heading = doc.add_heading(f"{exp.job_title} - {exp.company}", level=2)
                    
                    # Date range
                    date_range = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
                    if exp.location:
                        date_range += f" | {exp.location}"
                    
                    date_para = doc.add_paragraph(date_range)
                    date_para.runs[0].font.italic = True
                    date_para.runs[0].font.size = Pt(10)
                    
                    # Description
                    if exp.description:
                        doc.add_paragraph(exp.description)
                    
                    # Achievements
                    if exp.achievements:
                        for achievement in exp.achievements:
                            doc.add_paragraph(achievement, style='List Bullet')
            
            # Education
            if cv_data.education:
                doc.add_heading('Education', level=1)
                for edu in cv_data.education:
                    edu_heading = doc.add_heading(f"{edu.degree} - {edu.institution}", level=2)
                    
                    date_range = f"{edu.start_date or ''} - {edu.end_date or ''}"
                    if edu.location:
                        date_range += f" | {edu.location}"
                    
                    date_para = doc.add_paragraph(date_range)
                    date_para.runs[0].font.italic = True
                    date_para.runs[0].font.size = Pt(10)
                    
                    if edu.gpa:
                        doc.add_paragraph(f"GPA: {edu.gpa}")
            
            # Skills
            if cv_data.skills or cv_data.skill_categories:
                doc.add_heading('Skills', level=1)
                
                if cv_data.skill_categories:
                    for category in cv_data.skill_categories:
                        doc.add_paragraph(f"{category.category}: {', '.join(category.skills)}")
                elif cv_data.skills:
                    doc.add_paragraph(', '.join(cv_data.skills))
            
            # Certifications
            if cv_data.certifications:
                doc.add_heading('Certifications', level=1)
                for cert in cv_data.certifications:
                    cert_text = f"{cert.name} - {cert.issuer}"
                    if cert.issue_date:
                        cert_text += f" ({cert.issue_date})"
                    doc.add_paragraph(cert_text, style='List Bullet')
            
            # Projects
            if cv_data.projects:
                doc.add_heading('Projects', level=1)
                for project in cv_data.projects:
                    doc.add_heading(project.name, level=2)
                    doc.add_paragraph(project.description)
                    if project.technologies:
                        doc.add_paragraph(f"Technologies: {', '.join(project.technologies)}")
            
            # Languages
            if cv_data.languages:
                doc.add_heading('Languages', level=1)
                for lang in cv_data.languages:
                    doc.add_paragraph(f"{lang.language}: {lang.proficiency}", style='List Bullet')
            
            # Generate filename
            if output_filename is None:
                safe_name = cv_data.full_name.replace(' ', '_').replace('/', '_')
                output_filename = f"{safe_name}_CV_{uuid.uuid4().hex[:8]}.docx"
            
            if not output_filename.endswith('.docx'):
                output_filename += '.docx'
            
            output_path = self.output_dir / output_filename
            
            # Save document
            doc.save(str(output_path))
            
            logger.success(f"DOCX exported successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise Exception(f"Failed to export DOCX: {e}")
    
    async def export(
        self,
        cv_data: CVData,
        template_name: str = "modern_cv",
        output_format: Literal["pdf", "docx"] = "pdf",
        output_filename: str = None
    ) -> Path:
        """
        Export CV data to specified format.
        
        Args:
            cv_data: Structured CV data
            template_name: Template to use
            output_format: Output format (pdf or docx)
            output_filename: Custom output filename (optional)
            
        Returns:
            Path: Path to generated file
        """
        if output_format.lower() == "pdf":
            return await self.export_to_pdf(cv_data, template_name, output_filename)
        elif output_format.lower() == "docx":
            return await self.export_to_docx(cv_data, template_name, output_filename)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
