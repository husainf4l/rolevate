"""File parsing utilities for different document formats."""
import io
from pathlib import Path
from typing import Union
from loguru import logger

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


class FileParser:
    """Parse various file formats to extract text content."""
    
    async def parse_file(self, file_path: Union[str, Path]) -> str:
        """
        Parse a file and extract text content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            str: Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return await self._parse_pdf(file_path)
        elif ext == '.docx':
            return await self._parse_docx(file_path)
        elif ext in ['.txt', '.json']:
            return await self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    async def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF file."""
        if PdfReader is None:
            raise ImportError("PyPDF2 is required for PDF parsing")
        
        logger.info(f"Parsing PDF: {file_path}")
        
        try:
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            logger.success(f"Extracted {len(content)} characters from PDF")
            return content
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise Exception(f"Failed to parse PDF: {e}")
    
    async def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX file."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX parsing")
        
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            doc = Document(str(file_path))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            content = "\n".join(text_parts)
            logger.success(f"Extracted {len(content)} characters from DOCX")
            return content
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise Exception(f"Failed to parse DOCX: {e}")
    
    async def _parse_text(self, file_path: Path) -> str:
        """Parse plain text or JSON file."""
        logger.info(f"Reading text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            logger.success(f"Read {len(content)} characters from file")
            return content
            
        except Exception as e:
            logger.error(f"Error reading file: {e}")
            raise Exception(f"Failed to read file: {e}")
    
    async def parse_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Parse file from bytes.
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (for extension detection)
            
        Returns:
            str: Extracted text content
        """
        ext = Path(filename).suffix.lower()
        
        if ext == '.pdf':
            if PdfReader is None:
                raise ImportError("PyPDF2 is required for PDF parsing")
            
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
            return "\n\n".join(text_parts)
            
        elif ext == '.docx':
            if Document is None:
                raise ImportError("python-docx is required for DOCX parsing")
            
            doc = Document(io.BytesIO(file_bytes))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(text_parts)
            
        elif ext in ['.txt', '.json']:
            return file_bytes.decode('utf-8')
            
        else:
            raise ValueError(f"Unsupported file format: {ext}")
